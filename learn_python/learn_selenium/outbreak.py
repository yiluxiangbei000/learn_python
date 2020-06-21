__author__ = 'Administrator'

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import time
import datetime
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Pt

today = datetime.datetime.today()
yesterday = today - datetime.timedelta(days=1)
yr, mt, dy = today.year, today.month, today.day
today = '%s年%s月%s日'%(yr, mt, dy)
yr_1, mt_1, dy_1 = yesterday.year, yesterday.month, yesterday.day
yesterday = '%s年%s月%s日'%(yr_1, mt_1, dy_1)

def get_outbreak():
    option = Options()
    option.add_argument('--headless')
    driver = webdriver.Chrome(chrome_options=option)
    driver.get('https://www.outbreak.my/')
    time.sleep(3)

    #confirmed = driver.find_element_by_id('cases-my-confirmed').text
    confirmed = driver.find_element_by_xpath('//p[@class="value-digit"]/span[@id="cases-my-confirmed"]').text
    active = driver.find_element_by_id('cases-my-active').text
    death = driver.find_element_by_id('cases-my-death').text
    cofirmed_changes = driver.find_element_by_id('cases-my-confirmed-changes').text[2:]
    death_changes = driver.find_element_by_id('cases-my-death-changes').text

    states = driver.find_elements_by_xpath('//div[@class="card-body o-auto"]//tr')
    date = driver.find_element_by_xpath('//span[@id="last-update"]').text[:9]
    for state in states[1:]:
        if state.find_element_by_css_selector('td.text-value').text == 'Melaka':
            melaka_confirmed = state.find_element_by_css_selector('td.text-right.text-value-total').text
            melaka_death = state.find_element_by_css_selector('td.text-right.text-value-black').text
            try:
                melaka_changes = state.find_element_by_css_selector('td.fe.fe-arrow-up.float-center').text
            except:
                melaka_changes = 0
            break

    text = '疫情报告：截止到{}，马来西亚全国单日新增确诊{}例，累计确诊{}例，现存病例{}例，累计死亡{}例；项目部所在地马六甲州新增{}例，累计确诊{}例，累计死亡{}例。'.\
        format(yesterday, cofirmed_changes, confirmed, active, death, melaka_changes, melaka_confirmed, melaka_death )
    print(text)
    driver.quit()
    return text

def generate_docx(text):
    try:
	    doc = Document('D://_programming exercise//twentysixmay//learn_selenium//总承包分公司项目复工复产及防疫日报-马来项目%s0%s%s.docx' %(yr_1, mt_1, dy_1))
    except:
	    print('找不到昨天的日报，请核查')
    p = doc.paragraphs
    p7text, p10text, p_5text = p[7].text, p[10].text, p[-5].text
    p[7].text, p[10].text, p[-5].text = '', '', ''

    p[7].add_run(p7text.replace(yesterday, today)).font.size = Pt(16)
    p[10].add_run(text).font.size = Pt(16)
    p[-5].add_run(p_5text.replace(yesterday, today)).font.size = Pt(16)

    doc.save('D://_programming exercise//twentysixmay//learn_selenium//总承包分公司项目复工复产及防疫日报-马来项目%s0%s%s.docx' %(yr, mt, dy))
    print('总承包分公司项目复工复产及防疫日报-马来项目%s0%s%s.docx已生成！' %(yr, mt, dy))

if __name__ == '__main__':
    text = get_outbreak()
    generate_docx(text)







